
import json
from docx import Document
from datetime import datetime, timedelta

# 🧾 Constants
INPUT_DOCX_PATH = "working_agreement.docx"
OUTPUT_DOCX_PATH = "working_agreement.docx"
JSON_PATH = "extracted_values.json"
LABEL = "Maturity date is "
KEY = "maturity_date"

def parse_date(value):
    """Parse the input date from various formats to 'Month Day, Year'."""
    try:
        # Excel float date
        if isinstance(value, (int, float)):
            base_date = datetime(1899, 12, 30)
            return (base_date + timedelta(days=float(value))).strftime("%B %d, %Y")

        # String formats
        for fmt in ("%Y-%m-%d %H:%M:%S", "%m/%d/%Y", "%B %d, %Y", "%Y/%m/%d"):
            try:
                return datetime.strptime(value, fmt).strftime("%B %d, %Y")
            except:
                continue

        # ISO format fallback
        return datetime.fromisoformat(value).strftime("%B %d, %Y")
    except Exception as e:
        return None

def process_doc():
    print("\n📘 Running fill_docx_v49.py")
    print("📂 Loading document and JSON...")

    doc = Document(INPUT_DOCX_PATH)
    with open(JSON_PATH) as f:
        data = json.load(f)

    raw_value = data.get(KEY)
    if raw_value is None:
        print(f"⚠️ Warning: Key '{KEY}' not found in JSON.")
        return

    formatted_date = parse_date(raw_value)
    if not formatted_date:
        print(f"⚠️ Warning: Could not parse date from value: {raw_value}")
        return

    print(f"📅 Parsed date: {formatted_date}")

    found = False

    for para in doc.paragraphs:
        if para.text.strip().startswith(LABEL):
            print(f"🔍 Found paragraph starting with label: '{LABEL}'")
            found = True

            last_run = para.runs[-1] if para.runs else None
            font_name = last_run.font.name if last_run else None
            font_size = last_run.font.size if last_run else None

            print(f"📝 Font used - Name: {font_name}, Size: {font_size.pt if font_size else 'Default'}")

            para.clear()
            parts = [LABEL, formatted_date, "."]
            for text in parts:
                run = para.add_run(text)
                run.font.name = font_name
                run.font.size = font_size
            break

    if not found:
        print(f"⚠️ Warning: Paragraph starting with label '{LABEL}' not found.")

    doc.save(OUTPUT_DOCX_PATH)
    print(f"💾 Document saved as: {OUTPUT_DOCX_PATH}")

if __name__ == "__main__":
    process_doc()
    print("✅ Script finished.")
