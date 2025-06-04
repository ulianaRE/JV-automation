
import json
import re
from docx import Document
from docx.shared import Pt

# 🧾 Constants
INPUT_DOCX_PATH = "working_agreement.docx"
OUTPUT_DOCX_PATH = "working_agreement.docx"
JSON_PATH = "extracted_values.json"
TARGET_LABEL = "Party B"
INSERT_LABEL = "Phone:"
KEY = "funding_partner1_phone"

def normalize_us_phone(phone_raw):
    # Remove all non-digit characters
    digits = re.sub(r"\D", "", phone_raw)
    if len(digits) != 10:
        raise ValueError("Phone number must have exactly 10 digits after cleanup.")
    return f"({digits[:3]}) {digits[3:6]}-{digits[6:]}"

def process_doc():
    print("\n📘 Running fill_docx_v45.py")
    print("📂 Loading document and JSON...")

    # Load the document and JSON file
    doc = Document(INPUT_DOCX_PATH)
    with open(JSON_PATH) as f:
        data = json.load(f)

    # Extract and normalize the phone number
    raw_value = data.get(KEY)
    if not raw_value:
        print(f"❌ Value for key '{KEY}' not found in JSON.")
        return

    try:
        value = normalize_us_phone(raw_value)
        print(f"✅ Normalized phone number: '{value}'")
    except ValueError as e:
        print(f"❌ Error: {e}")
        return

    found_party_b = False
    found_phone_label = False

    # Search for insertion point
    for i, para in enumerate(doc.paragraphs):
        if not found_party_b and TARGET_LABEL.lower() in para.text.strip().lower():
            print(f"🔍 Found 'Party B' at paragraph {i}")
            found_party_b = True
            continue

        if found_party_b and INSERT_LABEL in para.text:
            print(f"🔍 Found 'Phone:' in paragraph {i}")
            found_phone_label = True

            # Preserve font formatting from original paragraph
            run = para.runs[0] if para.runs else para.add_run()
            font_name = run.font.name or "Arial"
            font_size = run.font.size.pt if run.font.size else 11
            print(f"🖋 Preserving font: {font_name}, {font_size} pt")

            # Replace paragraph content
            para.clear()
            run = para.add_run(f"{INSERT_LABEL} {value}")
            run.font.name = font_name
            run.font.size = Pt(font_size)
            break

    if not found_party_b:
        print("❌ 'Party B' label not found.")
    elif not found_phone_label:
        print("❌ 'Phone:' label not found after 'Party B'. No insertion made.")

    # Save the modified document
    doc.save(OUTPUT_DOCX_PATH)
    print(f"💾 Saved output as: {OUTPUT_DOCX_PATH}")
    print("✅ Done.")

if __name__ == "__main__":
    process_doc()
