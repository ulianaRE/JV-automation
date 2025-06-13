
import json
from docx import Document
from docx.shared import Pt

# 🧾 Constants
INPUT_DOCX_PATH = "working_agreement.docx"
OUTPUT_DOCX_PATH = "working_agreement.docx"
JSON_PATH = "extracted_values.json"
TARGET_LABEL = "Party B"
INSERT_LABEL = "Address:"
KEY = "funding_partner1_address"

def process_doc():
    print("\n📘 Running fill_docx_v43.py")
    print("📂 Loading document and JSON...")

    # Load the Word document and JSON file
    doc = Document(INPUT_DOCX_PATH)
    with open(JSON_PATH) as f:
        data = json.load(f)

    # Retrieve the value from the JSON data
    value = data.get(KEY)
    if not value:
        print(f"❌ Value for key '{KEY}' not found in JSON.")
        return
    print(f"✅ Value to insert: '{value}'")

    found_party_b = False
    found_address_label = False

    # Iterate over all paragraphs to find the correct insertion point
    for i, para in enumerate(doc.paragraphs):
        if not found_party_b and TARGET_LABEL.lower() in para.text.strip().lower():
            print(f"🔍 Found 'Party B' at paragraph {i}")
            found_party_b = True
            continue

        if found_party_b and INSERT_LABEL in para.text:
            print(f"🔍 Found 'Address:' in paragraph {i}")
            found_address_label = True

            # Extract the font name and size from the first run to preserve formatting
            run = para.runs[0] if para.runs else para.add_run()
            font_name = run.font.name or "Arial"
            font_size = run.font.size.pt if run.font.size else 11
            print(f"🖋 Preserving font: {font_name}, {font_size} pt")

            # Replace the paragraph text with label and value
            para.clear()

            # Bold label
            label_run = para.add_run(INSERT_LABEL)
            label_run.bold = True
            label_run.font.name = font_name
            label_run.font.size = Pt(font_size)

            # Space + non-bold value
            value_run = para.add_run(f" {value}")
            value_run.font.name = font_name
            value_run.font.size = Pt(font_size)

            break

    if not found_party_b:
        print("❌ 'Party B' label not found.")
    elif not found_address_label:
        print("❌ 'Address:' label not found after 'Party B'. No insertion made.")

    # Save the updated document
    doc.save(OUTPUT_DOCX_PATH)
    print(f"💾 Saved output as: {OUTPUT_DOCX_PATH}")
    print("✅ Done.")

if __name__ == "__main__":
    process_doc()
