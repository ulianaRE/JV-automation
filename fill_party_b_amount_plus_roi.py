
import json
from docx import Document

# 🧾 Constants
INPUT_DOCX_PATH = "working_agreement.docx"
OUTPUT_DOCX_PATH = "working_agreement.docx"
JSON_PATH = "extracted_values.json"
LABEL = "B is entitled to receive the return of their funding"
KEY_AMOUNT = "funding_partner1_funding"
KEY_FEE = "funding_partner1_ROI"
BALLOT_BOX = "☒"

def format_currency(value):
    """Formats a numeric value as currency, e.g., 1250000 → $1,250,000.00"""
    try:
        return "${:,.2f}".format(float(value))
    except ValueError:
        return value  # fallback in case of non-numeric

def process_doc():
    print("\n📘 Running fill_docx_v47.py")
    print("📂 Loading document and JSON...")

    # Load the Word document and JSON data
    doc = Document(INPUT_DOCX_PATH)
    with open(JSON_PATH) as f:
        data = json.load(f)

    # Retrieve required values from JSON
    raw_amount = data.get(KEY_AMOUNT)
    raw_fee = data.get(KEY_FEE)
    if not raw_amount or not raw_fee:
        print(f"⚠️ Warning: Required key(s) '{KEY_AMOUNT}' or '{KEY_FEE}' not found in JSON.")
        return

    formatted_amount = format_currency(raw_amount)
    formatted_fee = format_currency(raw_fee)

    print(f"💵 Funding amount: {formatted_amount}")
    print(f"💵 Lending fee: {formatted_fee}")

    found = False

    # Search for the paragraph that contains the label
    for para in doc.paragraphs:
        if LABEL in para.text:
            print(f"🔍 Found paragraph with label text: '{LABEL}'")
            found = True

            # Extract style from the first run of the paragraph
            first_run = para.runs[0] if para.runs else None
            font_name = first_run.font.name if first_run else None
            font_size = first_run.font.size if first_run else None
            print(f"📝 Font used - Name: {font_name}, Size: {font_size.pt if font_size else 'Default'}")

            # Replace the paragraph with the new structured sentence
            para.clear()
            parts = [
                BALLOT_BOX,
                " ",
                "Party B is entitled to receive the return of their funding contribution in the amount of ",
                formatted_amount,
                ", a lending fee of ",
                formatted_fee,
                ", in addition to extension fees."
            ]
            for text in parts:
                run = para.add_run(text)
                run.font.name = font_name
                run.font.size = font_size
                if text in [formatted_amount, formatted_fee]:
                    run.bold = True
            break

    if not found:
        print(f"⚠️ Warning: Paragraph with label '{LABEL}' not found.")

    # Save the updated document
    doc.save(OUTPUT_DOCX_PATH)
    print(f"💾 Document saved as: {OUTPUT_DOCX_PATH}")

if __name__ == "__main__":
    process_doc()
    print("✅ Script finished.")
