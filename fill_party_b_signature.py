import docx
import json

# 📌 Constants
TARGET_PARAGRAPH_START = "Party B LLC"
INPUT_DOCX_PATH = "working_agreement.docx"
OUTPUT_DOCX_PATH = "working_agreement.docx"
JSON_PATH = "extracted_values.json"

# U.S. state abbreviations to full names
state_names = {
    "AL": "Alabama", "AK": "Alaska", "AZ": "Arizona", "AR": "Arkansas", "CA": "California",
    "CO": "Colorado", "CT": "Connecticut", "DE": "Delaware", "FL": "Florida", "GA": "Georgia",
    "HI": "Hawaii", "ID": "Idaho", "IL": "Illinois", "IN": "Indiana", "IA": "Iowa",
    "KS": "Kansas", "KY": "Kentucky", "LA": "Louisiana", "ME": "Maine", "MD": "Maryland",
    "MA": "Massachusetts", "MI": "Michigan", "MN": "Minnesota", "MS": "Mississippi", "MO": "Missouri",
    "MT": "Montana", "NE": "Nebraska", "NV": "Nevada", "NH": "New Hampshire", "NJ": "New Jersey",
    "NM": "New Mexico", "NY": "New York", "NC": "North Carolina", "ND": "North Dakota",
    "OH": "Ohio", "OK": "Oklahoma", "OR": "Oregon", "PA": "Pennsylvania", "RI": "Rhode Island",
    "SC": "South Carolina", "SD": "South Dakota", "TN": "Tennessee", "TX": "Texas", "UT": "Utah",
    "VT": "Vermont", "VA": "Virginia", "WA": "Washington", "WV": "West Virginia", "WI": "Wisconsin",
    "WY": "Wyoming"
}

print("\n🔄 Running fill_party_b_signature.py")

# Load data from JSON file
with open(JSON_PATH, "r") as f:
    data = json.load(f)
    print(f"📥 Loaded JSON data from: {JSON_PATH}")

# Extract values
state_abbr = data.get("funding_partner1_state", "").strip()
entity = data.get("funding_partner1_entity", "").strip()
print(f"📦 Extracted entity: '{entity}'")
print(f"🗺️ Extracted state: '{state_abbr}'")

# Load the Word document
doc = docx.Document(INPUT_DOCX_PATH)
print(f"📄 Loaded Word document: {INPUT_DOCX_PATH}")

# Loop through paragraphs to find the one to modify
for para in doc.paragraphs:
    if para.text.strip().startswith(TARGET_PARAGRAPH_START):
        print(f"🔍 Found target paragraph starting with '{TARGET_PARAGRAPH_START}'")

        # Use formatting from the first run
        if not para.runs:
            print("❌ Paragraph has no runs to copy style from. Skipping.")
            continue

        run = para.runs[0]
        font_name = run.font.name
        font_size = run.font.size
        print(f"🎨 Detected style -> Font: {font_name}, Size: {font_size}")

        # Determine what to insert
        if state_abbr.lower() == "an individual":
            result = f"{entity}, an individual"
            print("👤 State is 'an individual' ➜ Inserting personal format.")
        elif not state_abbr:
            result = entity
            print("⚠️ State is empty ➜ Inserting entity only.")
        elif state_abbr not in state_names:
            result = f"{entity}, {state_abbr}"
            print(f"❓ Unrecognized state ➜ Inserting raw: '{result}'")
        else:
            full_state = state_names[state_abbr]
            article = "an" if full_state[0].lower() in "aeiou" else "a"
            clean_entity = entity.replace(",", "")
            result = f"{clean_entity}, {article} {full_state} limited liability company"
            print(f"🏛 Recognized state ➜ Inserting: '{result}'")

        # Clear old paragraph text
        para.clear()

        # Add new styled text
        new_run = para.add_run(result)
        new_run.font.name = font_name
        new_run.font.size = font_size

        print(f"✅ Paragraph updated with: '{result}'")
        break
else:
    print(f"❌ No paragraph found starting with '{TARGET_PARAGRAPH_START}'")

# Save updated document
doc.save(OUTPUT_DOCX_PATH)
print(f"💾 Document saved to: {OUTPUT_DOCX_PATH}")
