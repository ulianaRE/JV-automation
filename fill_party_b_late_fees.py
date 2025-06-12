
import json
from docx import Document

# 🧾 Constants
INPUT_DOCX_PATH = "working_agreement.docx"
OUTPUT_DOCX_PATH = "working_agreement.docx"
JSON_PATH = "extracted_values.json"
LABEL = "Party B will guarantee only one extension of up to one (1) month for a fee of "
KEY = "funding_partner1_late_fee"

def format_currency(value):
    """Formats a numeric value as currency, e.g., 1250000 → $1,250,000.00"""
    try:
        return "${:,.2f}".format(float(value))
    except ValueError:
        return value  # fallback in case of non-numeric

def process_doc():
    print("\n📘 Running fill_party_b_late_fees.py")
    print("📂 Loading document and JSON...")

    doc = Document(INPUT_DOCX_PATH)
    with open(JSON_PATH) as f:
        data = json.load(f)

    raw_value = data.get(KEY)
    if raw_value is None:
        print(f"⚠️ Warning: Key '{KEY}' not found in JSON.")
        return

    formatted_value = format_currency(raw_value)
    print(f"💵 Late fee value to insert: {formatted_value}")

    found = False

    for para in doc.paragraphs:
        if para.text.strip().startswith(LABEL):
            print(f"🔍 Found paragraph starting with label: '{LABEL}'")
            found = True

            first_run = para.runs[0] if para.runs else None
            font_name = first_run.font.name if first_run else None
            font_size = first_run.font.size if first_run else None

            print(f"📝 Font used - Name: {font_name}, Size: {font_size.pt if font_size else 'Default'}")

            para.clear()

            # Add label part
            run_label = para.add_run(LABEL)
            run_label.font.name = font_name
            run_label.font.size = font_size

            # Add bold value
            run_value = para.add_run(formatted_value)
            run_value.bold = True
            run_value.font.name = font_name
            run_value.font.size = font_size

            # Add closing period
            run_period = para.add_run(".")
            run_period.font.name = font_name
            run_period.font.size = font_size

            break

    if not found:
        print(f"⚠️ Warning: Paragraph starting with label '{LABEL}' not found.")

    doc.save(OUTPUT_DOCX_PATH)
    print(f"💾 Document saved as: {OUTPUT_DOCX_PATH}")

if __name__ == "__main__":
    process_doc()
    print("✅ Script finished.")
