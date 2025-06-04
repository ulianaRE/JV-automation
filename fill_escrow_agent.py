
import json
from docx import Document
from docx.shared import Pt

# 🧾 Constants
INPUT_DOCX_PATH = "working_agreement.docx"
OUTPUT_DOCX_PATH = "working_agreement.docx"
JSON_PATH = "extracted_values.json"
KEY_NAME = "title_company_name"
KEY_EMAIL = "title_company_email"
PRIMARY_LABEL = "1.4 Closing Date of Transaction:"
TARGET_LABEL = "Escrow Agent:"

def process_doc():
    print("\n📘 Running fill_docx_v50.py")
    print("📂 Loading document and JSON...")

    # Load the document and the JSON file
    doc = Document(INPUT_DOCX_PATH)
    with open(JSON_PATH) as f:
        data = json.load(f)

    # Extract both values
    name = data.get(KEY_NAME)
    email = data.get(KEY_EMAIL)
    if not name or not email:
        print(f"⚠️ Missing value for key '{KEY_NAME}' or '{KEY_EMAIL}' in JSON. Skipping.")
        return

    value = f"{name} ({email})"
    print(f"📌 Combined value to insert: {value}")

    found_primary = False
    found_target = False

    # Locate primary and target labels
    for i, para in enumerate(doc.paragraphs):
        if not found_primary and PRIMARY_LABEL.lower() in para.text.strip().lower():
            print(f"🔍 Found '{PRIMARY_LABEL}' at paragraph {i}")
            found_primary = True
            continue

        if found_primary and TARGET_LABEL in para.text:
            print(f"🔍 Found '{TARGET_LABEL}' at paragraph {i}")
            found_target = True

            # Preserve formatting
            run = para.runs[0] if para.runs else para.add_run()
            font_name = run.font.name or "Arial"
            font_size = run.font.size.pt if run.font.size else 11
            print(f"🖋 Preserving font: {font_name}, {font_size} pt")

            # Replace content
            para.clear()
            run = para.add_run(f"{TARGET_LABEL} {value}")
            run.font.name = font_name
            run.font.size = Pt(font_size)
            break

    if not found_primary:
        print(f"⚠️ Label '{PRIMARY_LABEL}' not found. Skipping.")
    elif not found_target:
        print(f"⚠️ Label '{TARGET_LABEL}' not found after '{PRIMARY_LABEL}'. No insertion made.")

    # Save updated document
    doc.save(OUTPUT_DOCX_PATH)
    print(f"💾 Saved document as: {OUTPUT_DOCX_PATH}")
    print("✅ Done.")

if __name__ == "__main__":
    process_doc()
