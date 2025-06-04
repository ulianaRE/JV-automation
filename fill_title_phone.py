
import json
import re
from docx import Document
from docx.shared import Pt

# 🧾 Constants
INPUT_DOCX_PATH = "working_agreement.docx"
OUTPUT_DOCX_PATH = "working_agreement.docx"
JSON_PATH = "extracted_values.json"
KEY = "title_company_phone"
PRIMARY_LABEL = "1.4 Closing Date of Transaction:"
TARGET_LABEL = "Phone:"

def normalize_us_phone(phone_raw):
    """Normalize various phone formats to U.S. standard (xxx) xxx-xxxx."""
    digits = re.sub(r"\D", "", phone_raw)
    if len(digits) != 10:
        raise ValueError("Phone number must have exactly 10 digits after cleanup.")
    return f"({digits[:3]}) {digits[3:6]}-{digits[6:]}"

def process_doc():
    print("\n📘 Running fill_docx_v49.py")
    print("📂 Loading document and JSON...")

    # Load the document and JSON file
    doc = Document(INPUT_DOCX_PATH)
    with open(JSON_PATH) as f:
        data = json.load(f)

    # Extract and normalize the phone number
    raw_value = data.get(KEY)
    if not raw_value:
        print(f"⚠️ Value for key '{KEY}' not found in JSON. Skipping.")
        return

    try:
        value = normalize_us_phone(raw_value)
        print(f"✅ Normalized phone number: '{value}'")
    except ValueError as e:
        print(f"⚠️ Error: {e}. Skipping.")
        return

    found_primary = False
    found_target = False

    # Locate the required markers
    for i, para in enumerate(doc.paragraphs):
        if not found_primary and PRIMARY_LABEL.lower() in para.text.strip().lower():
            print(f"🔍 Found '{PRIMARY_LABEL}' at paragraph {i}")
            found_primary = True
            continue

        if found_primary and TARGET_LABEL in para.text:
            print(f"🔍 Found '{TARGET_LABEL}' at paragraph {i}")
            found_target = True

            # Preserve formatting
            run = para.runs[0] if para.runs else para.add_run()
            font_name = run.font.name or "Arial"
            font_size = run.font.size.pt if run.font.size else 11
            print(f"🖋 Preserving font: {font_name}, {font_size} pt")

            # Replace content
            para.clear()
            run = para.add_run(f"{TARGET_LABEL} {value}")
            run.font.name = font_name
            run.font.size = Pt(font_size)
            break

    if not found_primary:
        print(f"⚠️ Label '{PRIMARY_LABEL}' not found. Skipping.")
    elif not found_target:
        print(f"⚠️ Label '{TARGET_LABEL}' not found after '{PRIMARY_LABEL}'. No insertion made.")

    # Save modified document
    doc.save(OUTPUT_DOCX_PATH)
    print(f"💾 Saved document as: {OUTPUT_DOCX_PATH}")
    print("✅ Done.")

if __name__ == "__main__":
    process_doc()
