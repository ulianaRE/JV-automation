
import json
from docx import Document
from docx.shared import Pt

# 🧾 Constants
INPUT_DOCX_PATH = "working_agreement.docx"
OUTPUT_DOCX_PATH = "working_agreement.docx"
JSON_PATH = "extracted_values.json"
KEY = "title_company_entity"
PRIMARY_LABEL = "1.4 Closing Date of Transaction:"
TARGET_LABEL = "Title Company:"

def process_doc():
    print("\n📘 Running fill_docx_v47.py")
    print("📂 Loading document and JSON...")

    # Load the Word document and JSON values
    doc = Document(INPUT_DOCX_PATH)
    with open(JSON_PATH) as f:
        data = json.load(f)

    # Get the target value from JSON
    value = data.get(KEY)
    if not value:
        print(f"⚠️ Value for key '{KEY}' not found in JSON. Skipping.")
        return
    print(f"✅ Value to insert: '{value}'")

    found_primary = False
    found_target = False

    # Search for PRIMARY_LABEL and then TARGET_LABEL
    for i, para in enumerate(doc.paragraphs):
        if not found_primary and PRIMARY_LABEL.lower() in para.text.strip().lower():
            print(f"🔍 Found '{PRIMARY_LABEL}' at paragraph {i}")
            found_primary = True
            continue

        if found_primary and TARGET_LABEL in para.text:
            print(f"🔍 Found '{TARGET_LABEL}' at paragraph {i}")
            found_target = True

            # Get font name and size from first run
            run = para.runs[0] if para.runs else para.add_run()
            font_name = run.font.name or "Arial"
            font_size = run.font.size.pt if run.font.size else 11
            print(f"🖋 Preserving font: {font_name}, {font_size} pt")

            # Replace the paragraph content
            para.clear()

            # Bold label
            label_run = para.add_run(TARGET_LABEL)
            label_run.bold = True
            label_run.font.name = font_name
            label_run.font.size = Pt(font_size)

            # Space + regular value
            value_run = para.add_run(f" {value}")
            value_run.font.name = font_name
            value_run.font.size = Pt(font_size)

            break

    if not found_primary:
        print(f"⚠️ Label '{PRIMARY_LABEL}' not found. Skipping.")
    elif not found_target:
        print(f"⚠️ Label '{TARGET_LABEL}' not found after '{PRIMARY_LABEL}'. No insertion made.")

    # Save the document
    doc.save(OUTPUT_DOCX_PATH)
    print(f"💾 Saved output as: {OUTPUT_DOCX_PATH}")
    print("✅ Done.")

if __name__ == "__main__":
    process_doc()
