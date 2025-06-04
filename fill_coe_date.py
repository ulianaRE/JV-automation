
import json
from docx import Document
from docx.shared import Pt
from datetime import datetime, timedelta
import re

# 🧾 Constants
INPUT_DOCX_PATH = "working_agreement.docx"
OUTPUT_DOCX_PATH = "working_agreement.docx"
JSON_PATH = "extracted_values.json"
KEY = "coe_date"
LABEL_TEXT = "1.4 Closing Date of Transaction:"

def parse_date(value):
    """Try to parse the date input into a datetime object."""
    try:
        if isinstance(value, (int, float)):
            # Excel timestamp (days since 1899-12-30)
            return datetime(1899, 12, 30) + timedelta(days=float(value))
        if isinstance(value, str):
            # Try multiple date formats
            for fmt in ("%Y-%m-%d %H:%M:%S", "%Y/%m/%d", "%m/%d/%Y", "%B %d, %Y", "%Y-%m-%d"):
                try:
                    return datetime.strptime(value, fmt)
                except ValueError:
                    continue
            # Fallback to dateutil.parser if available
            try:
                from dateutil.parser import parse
                return parse(value)
            except Exception:
                pass
    except Exception as e:
        print(f"⚠️ Error parsing date: {e}")
    return None

def process_doc():
    print("\n📘 Running fill_docx_v46.py")
    print("📂 Loading document and JSON...")

    doc = Document(INPUT_DOCX_PATH)
    with open(JSON_PATH) as f:
        data = json.load(f)

    raw_value = data.get(KEY)
    if not raw_value:
        print(f"⚠️ '{KEY}' not found in JSON. Skipping insertion.")
        return

    date_obj = parse_date(raw_value)
    if not date_obj:
        print(f"⚠️ Could not parse '{raw_value}' as a valid date. Skipping.")
        return

    formatted_date = date_obj.strftime("%B %d, %Y")
    print(f"✅ Parsed and formatted date: {formatted_date}")

    found_label = False
    for i, para in enumerate(doc.paragraphs):
        if LABEL_TEXT.lower() in para.text.strip().lower():
            print(f"🔍 Found target label in paragraph {i}")
            found_label = True

            # Preserve formatting from the first run
            run = para.runs[0] if para.runs else para.add_run()
            font_name = run.font.name or "Arial"
            font_size = run.font.size.pt if run.font.size else 11

            print(f"🖋 Preserving font: {font_name}, {font_size} pt")

            # Replace text content
            para.clear()
            new_text = f"{LABEL_TEXT} {formatted_date}"
            run = para.add_run(new_text)
            run.font.name = font_name
            run.font.size = Pt(font_size)

            break

    if not found_label:
        print(f"⚠️ Label '{LABEL_TEXT}' not found in document. No insertion made.")

    doc.save(OUTPUT_DOCX_PATH)
    print(f"💾 Saved document as: {OUTPUT_DOCX_PATH}")
    print("✅ Done.")

if __name__ == "__main__":
    process_doc()
