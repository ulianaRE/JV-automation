
import json
from docx import Document
from docx.shared import Pt

# 🧾 Constants
INPUT_DOCX_PATH = "working_agreement.docx"
OUTPUT_DOCX_PATH = "working_agreement.docx"
JSON_PATH = "extracted_values.json"
LABEL = "1.1 Property:"
KEY = "property_address"

def process_doc():
    print("\n📘 Running fill_property.py")
    print("📂 Loading document and JSON...")

    # Load the document and JSON data
    doc = Document(INPUT_DOCX_PATH)
    with open(JSON_PATH) as f:
        data = json.load(f)

    # Retrieve the value for insertion
    value = data.get(KEY)
    if not value:
        print(f"❌ Value for key '{KEY}' not found in JSON.")
        return

    print(f"✅ Value to insert: '{value}'")

    # Search for the paragraph with the label
    found = False
    for para in doc.paragraphs:
        if para.text.strip().startswith(LABEL):
            print(f"🔍 Found label: '{LABEL}'")
            found = True

            # Preserve original font settings
            run = para.runs[0]
            font_name = run.font.name or "Arial"
            font_size = run.font.size.pt if run.font.size else 11
            print(f"🖋 Preserving font: {font_name}, {font_size} pt")

            # Clear paragraph and insert new text with tab
            para.clear()
            run = para.add_run(LABEL + "\t" + value)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            break

    if not found:
        print(f"❌ Label '{LABEL}' not found. No insertion made.")

    # Save the modified document
    doc.save(OUTPUT_DOCX_PATH)
    print(f"💾 Saved output as: {OUTPUT_DOCX_PATH}")

if __name__ == "__main__":
    process_doc()
