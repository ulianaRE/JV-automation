#!/usr/bin/env python3
"""
fill_10_percent_v3.py

Loads a Word document and JSON data, computes 10% of owner_partner_funding,
and replaces the first $ placeholder after a specific label in the doc
with a bold, styled currency value inserted *in place* of the placeholder.
It rebuilds the paragraph runs to maintain correct order and formatting.
Outputs a new versioned document.
"""

import json
from docx import Document
from docx.shared import Pt

# 🧭 Constants
INPUT_DOCX = "working_agreement.docx"
INPUT_JSON = "extracted_values.json"
OUTPUT_DOCX = "working_agreement.docx"
SCRIPT_VERSION = 3
# for debugging only OUTPUT_DOCX = f"working_agreement_v{SCRIPT_VERSION}.docx"

LABEL = "Remedies on Continued Default:"
JSON_KEY = "owner_partner_funding"
PLACEHOLDER = "$"

def format_currency(value):
    """Format value as currency with two decimals."""
    return f"${value:,.2f}"

def main():
    print("\nStarting fill_10_percent_v3 script...\n")

    # Load JSON
    try:
        with open(INPUT_JSON, "r") as f:
            data = json.load(f)
    except Exception as e:
        print(f"ERROR loading JSON file: {e}")
        return

    if JSON_KEY not in data:
        print(f"WARNING: '{JSON_KEY}' not found. Exiting.")
        return

    try:
        funding = float(data[JSON_KEY])
    except ValueError:
        print(f"WARNING: '{JSON_KEY}' not numeric. Exiting.")
        return

    insert_text = format_currency(funding * 0.10)

    # Load Word document
    try:
        doc = Document(INPUT_DOCX)
    except Exception as e:
        print(f"ERROR loading docx: {e}")
        return

    target_para = next(
        (p for p in doc.paragraphs if p.text.strip().startswith(LABEL)),
        None
    )
    if not target_para:
        print(f"WARNING: Label '{LABEL}' paragraph not found. Exiting.")
        return

    print(f"Found paragraph: '{target_para.text[:60]}...'")

    # Rebuild runs
    new_runs = []
    replaced = False

    for run in target_para.runs:
        text = run.text
        if not replaced and PLACEHOLDER in text:
            before, after = text.split(PLACEHOLDER, 1)
            font_name = run.font.name
            font_size = run.font.size.pt if run.font.size else None

            # Text before $
            new_runs.append((before, run.bold, font_name, font_size))
            # Inserted bold value
            new_runs.append((insert_text, True, font_name, font_size))
            # Text after $
            new_runs.append((after, run.bold, font_name, font_size))

            replaced = True
        else:
            # Keep run unchanged
            new_runs.append((text, run.bold, run.font.name, 
                             run.font.size.pt if run.font.size else None))

    if not replaced:
        print("WARNING: '$' placeholder not found. Exiting.")
        return

    # Clear old runs
    for r in target_para.runs:
        r.clear()
    # Add rebuilt runs
    for text, bold, fname, fsize in new_runs:
        r = target_para.add_run(text)
        r.bold = bold
        if fname:
            r.font.name = fname
        if fsize:
            r.font.size = Pt(fsize)

    # Save and close
    doc.save(OUTPUT_DOCX)
    print(f"\nInserted '{insert_text}'.\nSaved as '{OUTPUT_DOCX}'.\nScript complete.\n")

if __name__ == "__main__":
    main()
