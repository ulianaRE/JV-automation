
import json
from docx import Document

# üßæ Constants
INPUT_DOCX_PATH = "working_agreement.docx"
OUTPUT_DOCX_PATH = "working_agreement.docx"
JSON_PATH = "extracted_values.json"
LABEL = "From the total funding contribution outlined in Section 1.5, an immediate sum of "
KEY = "owner_partner_funding"

def format_currency(value):
    """Formats a numeric value as currency, e.g., 1250000 ‚Üí $1,250,000.00"""
    try:
        return "${:,.2f}".format(float(value))
    except ValueError:
        return value  # fallback in case of non-numeric

def process_doc():
    print("\nüìò Running fill_funds_released_at_coe.py")
    print("üìÇ Loading document and JSON...")

    # Load the Word document and JSON data
    doc = Document(INPUT_DOCX_PATH)
    with open(JSON_PATH) as f:
        data = json.load(f)

    # Retrieve the value from JSON
    raw_value = data.get(KEY)
    if not raw_value:
        print(f"‚ö†Ô∏è Warning: Value for key '{KEY}' not found in JSON.")
        return

    formatted_value = format_currency(raw_value)
    print(f"üíµ Value to insert: {formatted_value}")

    found = False

    # Search for the paragraph that starts with the label
    for para in doc.paragraphs:
        if para.text.strip().startswith(LABEL):
            print(f"üîç Found paragraph starting with label: '{LABEL}'")
            found = True

            # Extract font/style from the 1st run
            first_run = para.runs[0] if para.runs else None
            font_name = first_run.font.name if first_run else None
            font_size = first_run.font.size if first_run else None

            print(f"üìù Font used - Name: {font_name}, Size: {font_size.pt if font_size else 'Default'}")

            # Clear the paragraph and reconstruct it
            para.clear()
            parts = [
                LABEL,
                formatted_value,
                " shall be designated and made promptly available. This allocation covers expenses such as property acquisition, closing costs, holding costs, insurance, transaction coordinator (TC) fees, and construction/renovations, as detailed in Section 1.1"
            ]
            for text in parts:
                run = para.add_run(text)
                run.font.name = font_name
                run.font.size = font_size
                run.bold = True if text == formatted_value else None
            break

    if not found:
        print(f"‚ö†Ô∏è Warning: Paragraph starting with label '{LABEL}' not found.")

    # Save the updated document
    doc.save(OUTPUT_DOCX_PATH)
    print(f"üíæ Document saved as: {OUTPUT_DOCX_PATH}")

if __name__ == "__main__":
    process_doc()
    print("‚úÖ Script finished.")
