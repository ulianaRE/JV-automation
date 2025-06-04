
import json
from docx import Document

# 🧾 Constants
INPUT_DOCX_PATH = "working_agreement.docx"
OUTPUT_DOCX_PATH = "working_agreement.docx"
JSON_PATH = "extracted_values.json"
LABEL = "capital investment of Party B"
KEY = "owner_partner_funding"
BALLOT_BOX = "☐"

def format_currency(value):
    """Formats a numeric value as currency, e.g., 1250000 → $1,250,000.00"""
    try:
        return "${:,.2f}".format(float(value))
    except ValueError:
        return value  # fallback in case of non-numeric

def process_doc():
    print("\n📘 Running fill_party_a_funding.py")
    print("📂 Loading document and JSON...")

    # Load the Word document and JSON data
    doc = Document(INPUT_DOCX_PATH)
    with open(JSON_PATH) as f:
        data = json.load(f)

    # Retrieve the value from JSON using the specified key
    raw_value = data.get(KEY)
    if not raw_value:
        print(f"⚠️ Warning: Value for key '{KEY}' not found in JSON.")
        return
    formatted_value = format_currency(raw_value)
    print(f"💵 Value to insert: {formatted_value}")

    found = False

    # Iterate over paragraphs to find the one starting with the label
    for para in doc.paragraphs:
        if LABEL in para.text:
            print(f"🔍 Found paragraph with label text: '{LABEL}'")
            found = True

            # Attempt to capture the style from the run that contains the label
            style = None
            for run in para.runs:
                if LABEL in run.text:
                    style = {
                        'font_name': run.font.name,
                        'font_size': run.font.size,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline
                    }
                    break

            # Fallback style if none found
            if not style:
                print("⚠️ Warning: Could not detect style from label run.")
                style = {
                    'font_name': None,
                    'font_size': None,
                    'bold': None,
                    'italic': None,
                    'underline': None
                }

            print(f"📝 Font used - Name: {style['font_name']}, Size: {style['font_size'].pt if style['font_size'] else 'Default'}")

            # Replace the paragraph with the new structured sentence
            para.clear()
            parts = [
                BALLOT_BOX,
                "  ",  # 2 spaces
                "Party A funding: ",
                "all amounts exceeding the ",
                formatted_value,
                " ",
                LABEL
            ]
            for text in parts:
                run = para.add_run(text)
                run.font.name = style['font_name']
                run.font.size = style['font_size']
                run.bold = style['bold']
                run.italic = style['italic']
                run.underline = style['underline']
            break

    if not found:
        print(f"⚠️ Warning: Paragraph with label '{LABEL}' not found.")

    # Save the modified document
    doc.save(OUTPUT_DOCX_PATH)
    print(f"💾 Document saved as: {OUTPUT_DOCX_PATH}")

if __name__ == "__main__":
    process_doc()
    print("✅ Script finished.")
