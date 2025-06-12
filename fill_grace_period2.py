
import json
from docx import Document
from num2words import num2words

# 🧾 Constants
INPUT_DOCX_PATH = "working_agreement.docx"
OUTPUT_DOCX_PATH = "working_agreement.docx"
JSON_PATH = "extracted_values.json"
LABEL_SNIPPET = "calendar day grace period"
KEY = "grace_period"

def process_doc():
    print("\n📘 Running fill_docx_v53.py")
    print("📂 Loading document and JSON...")

    doc = Document(INPUT_DOCX_PATH)
    with open(JSON_PATH) as f:
        data = json.load(f)

    # Extract grace period
    raw_value = data.get(KEY)
    if raw_value is None:
        print(f"⚠️ Warning: Key '{KEY}' not found in JSON.")
        return

    try:
        days = int(raw_value)
        word_form = num2words(days).lower()
    except Exception as e:
        print(f"❌ Error parsing grace period value '{raw_value}':", str(e))
        return

    formatted_string = (
        "A {} ({}) calendar day grace period will be provided after the Maturity Date, "
        "during which no extension/late fee will be charged. If the loan is not repaid in full within the grace period, "
        "the extension fee will begin accruing on day 2 and be calculated from that day forward."
    ).format(word_form, days)

    print(f"📆 Final formatted sentence: {formatted_string}")

    found = False

    for para in doc.paragraphs:
        if LABEL_SNIPPET in para.text:
            print(f"🔍 Found paragraph containing: '{LABEL_SNIPPET}'")
            found = True

            last_run = para.runs[-1] if para.runs else None
            font_name = last_run.font.name if last_run else None
            font_size = last_run.font.size if last_run else None

            print(f"📝 Font used - Name: {font_name}, Size: {font_size.pt if font_size else 'Default'}")

            # Replace paragraph text
            para.clear()
            run = para.add_run(formatted_string)
            run.font.name = font_name
            run.font.size = font_size
            break

    if not found:
        print(f"⚠️ Warning: Paragraph containing '{LABEL_SNIPPET}' not found.")

    doc.save(OUTPUT_DOCX_PATH)
    print(f"💾 Document saved as: {OUTPUT_DOCX_PATH}")

if __name__ == "__main__":
    process_doc()
    print("✅ Script finished.")
