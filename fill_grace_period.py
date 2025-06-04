
import json
from docx import Document

# 🧾 Constants
INPUT_DOCX_PATH = "working_agreement.docx"
OUTPUT_DOCX_PATH = "working_agreement.docx"
JSON_PATH = "extracted_values.json"
LABEL = "Grace period of "
KEY = "grace_period"

def process_doc():
    print("\n📘 Running fill_docx_v50.py")
    print("📂 Loading document and JSON...")

    # Load Word doc and JSON
    doc = Document(INPUT_DOCX_PATH)
    with open(JSON_PATH) as f:
        data = json.load(f)

    # Extract grace period value
    value = data.get(KEY)
    if value is None:
        print(f"⚠️ Warning: Key '{KEY}' not found in JSON.")
        return

    print(f"📆 Grace period value to insert: {value}")

    found = False

    # Search for paragraph with label
    for para in doc.paragraphs:
        if para.text.strip().startswith(LABEL):
            print(f"🔍 Found paragraph starting with label: '{LABEL}'")
            found = True

            # Extract font from last run
            last_run = para.runs[-1] if para.runs else None
            font_name = last_run.font.name if last_run else None
            font_size = last_run.font.size if last_run else None

            print(f"📝 Font used - Name: {font_name}, Size: {font_size.pt if font_size else 'Default'}")

            # Replace paragraph with constructed sentence
            para.clear()
            parts = [LABEL, str(value), " calendar days."]
            for text in parts:
                run = para.add_run(text)
                run.font.name = font_name
                run.font.size = font_size
            break

    if not found:
        print(f"⚠️ Warning: Paragraph starting with label '{LABEL}' not found.")

    doc.save(OUTPUT_DOCX_PATH)
    print(f"💾 Document saved as: {OUTPUT_DOCX_PATH}")

if __name__ == "__main__":
    process_doc()
    print("✅ Script finished.")
